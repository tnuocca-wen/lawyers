from openai import OpenAI
import os, json
import markdown2
import pymupdf
from docx import Document
from PIL import Image
from itertools import repeat
from concurrent.futures import ProcessPoolExecutor

from pathlib import Path
BASE_DIR = Path(__file__).resolve().parent.parent 

from django.conf import settings

import google.generativeai as genai

genai.configure(api_key=os.getenv("GOOGLE_GEN_API_KEY"))

client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))


def gpt_response(messages, json=False, temperature=0.5):
    # print(messages)
    response = client.chat.completions.create(
    response_format={ "type": "json_object" } if json else {"type": "text"}, #"json_schema", "json_schema": {'choice': 'the best option from the given choices' } } if json else {},
    temperature = temperature,
    model="gpt-4o-mini",
    messages=messages
    )

    return response.choices[0].message.content

def intro_description(content):
    messages = [
        {'role': 'system', 'content': '''You are a junior legal professional assisting a senior counsel.
You prepare notes for the senior counsel, after reviewing all documents. You never add any chatty texts in your responses.'''}]
    
    messages.append({"role": "user", "content": f"""
The following raw text was extracted from a scanned legal document using OCR (Optical Character Recognition) software.
Please review it carefully.

---------------------------------------
Raw Text:
{content}
---------------------------------------

Based on the above raw text, come up with a description of the above document,
without leaving out any information. Keep the description short and to the point.
Do not include generic information which a senior counsel would already know.
"""})
    return gpt_response(messages, False)


def notes(page, messages):
    
    messages.append({'role': 'user', 'content': f"""
The following is a page from a legal document.
---------------------------------------
**Document:**
{page}
---------------------------------------

**Your task:**
- Carefully review the above document.
- Prepare a note on the above document in bullet point format.
- Ensure that the note is accurate, and coherent."""})

    return gpt_response(messages, False)

def timeline(page, messages):
    messages.append({'role': 'user', 'content': f"""
The following is a page from a legal document.
---------------------------------------
**Document:**
{page}
---------------------------------------

**Your task:**
- Carefully review the above document.
- Prepare a chronology based on the dates mentioned in the document.
- Ensure that only those **information with a date** is returned
- Give no preface or explanation"""})
    
    return gpt_response(messages, False)

def fluidate(content, messages, choice):

    content = "\n\n\n".join(content)
    if choice == 0:
        messages.append({'role': 'user', 'content': f"""The following document is notes prepared based on individual pages of a legal document.
Since notes were prepared on each page and finally concatenated together, it is likely
that there is some lack of organization.
---------------------------------------
**Document:**
{content}
---------------------------------------

**Your task:**
- Carefully review the above document.
- Streamline the above document by removing repetitions.
- Organize the paragraphs logically"""})
    else:
        messages.append({'role': 'user', 'content': f"""The following document is timeline prepared based on individual pages of a legal document.
Since the timeline was prepared from each page and finally concatenated together, it is likely
that there is some lack of organization.
---------------------------------------
**Document:**
{content}
---------------------------------------

**Your task:**
- Carefully review the above document.
- Streamline the above document by removing repetitions.
- Organize the paragraphs logically"""})

    return gpt_response(messages, False)

def create_timeline(timelineList):

    timeline_content = "\n".join(timelineList)

    messages = [
        {'role': 'system', 'content': '''You are a junior legal professional assisting a senior counsel.
You prepare timelines for the senior counsel, after reviewing all documents. You never add any chatty texts in your responses.'''}]
    
    messages.append({"role": "user", "content": f"""
The following content was extracted multiple documents.
Please review it carefully.

---------------------------------------
Content to make timeline with:
{timeline_content}
---------------------------------------

Based on the above content, come up with the correct timeline, without leaving out any information.
"""})
    
    return gpt_response(messages, False)



def sk_wrapper(choice, fc):

    messages = [
        {'role': 'system', 'content': '''You are a junior legal professional assisting a senior counsel.
You prepare notes for the senior counsel, after reviewing all documents. You never add any chatty texts in your responses.'''}]

    output = None

    if choice == 0:
        with ProcessPoolExecutor() as executor:
            output = list(executor.map(notes, fc, repeat(messages)))
    elif choice == 1:
        with ProcessPoolExecutor() as executor:
            output = list(executor.map(timeline, fc, repeat(messages)))
    else:
        content = intro_description("".join(fc)[:3000])

    if output:
        content = fluidate(output, messages, choice)
    
    return content

def process(fl):
    print(fl)
    fl = fl.replace("\\", "/")
    if fl:
        if fl.split("/")[-1].split(".")[-1] =='txt':
            fc = extract_text_from_txt(fl)

        elif fl.split("/")[-1].split(".")[-1] =='pdf':
            print("Got here inside the PDF reading logic")
            fc = extract_text_from_pdf(fl)
        
        elif fl.split("/")[-1].split(".")[-1] =='docx':
            fc = extract_text_from_docx(fl)

        fc = make_chunks(fc, 2000, 900)
        
        dif_outs = [0, 1, 2]

        # sk = sk_wrapper(dif_outs, fc)

        with ProcessPoolExecutor() as executor:
            sk = list(executor.map(sk_wrapper, dif_outs, repeat(fc)))
        
        description = sk[2]
        
        sk = {"notes": sk[0], "timeline": sk[1]}
        

        # print(sk)
        return description, sk, fl

def make_chunks(texts, chunk_chars=2000, look_bfr=700):

    if chunk_chars <= 0 or look_bfr < 0:
        raise ValueError("chunk_chars must be positive and look_bfr must be non-negative.")
    if look_bfr >= chunk_chars:
        raise ValueError("look_bfr must be smaller than chunk_chars.")

    chunks = []
    start = 0
    text_length = len(texts)

    while start < text_length:
        end = min(start + chunk_chars, text_length)
        chunk = texts[start:end]

        # Append previous look-back context if needed
        if start != 0:
            overlap_start = max(0, start - look_bfr)
            chunk = texts[overlap_start:start] + chunk

        chunks.append(chunk)
        start += chunk_chars

    return chunks



def convert_markdown_to_html(markdown_text):
    html = markdown2.markdown(markdown_text)
    return html


def extract_text_from_txt(txt_path):
    with open(txt_path, 'r', encoding='utf-8') as file:
                content = file.read()
    return content

def extract_text_from_docx(doc_path):
    doc = Document(doc_path)
    content = []
    
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            content.append(f"### {paragraph.text}")
        else:
            content.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            content.append('\t'.join(row_data))

    return '\n'.join(content)


def extract_text_from_pdf(pdf_path):
    doc = pymupdf.open(pdf_path)
    zoom_x = 2.0
    zoom_y = 2.0
    mat = pymupdf.Matrix(zoom_x, zoom_y)
    doc_pics = []
    pdf_path = pdf_path.replace('\\', '/')
    print(pdf_path)
    bfp = os.path.join(settings.MEDIA_ROOT, f'file_images/{pdf_path.split("/")[-1].replace(".pdf", "")}').replace("\\", "/")
    print(bfp)
    if not os.path.isdir(bfp):
        os.mkdir(bfp)
    for page in doc:
        pix = page.get_pixmap(matrix=mat)
        fp = os.path.join(bfp, f"page-{page.number}.png")
        pix.save(fp)
        doc_pics.append(fp)
    # for path in doc_pics:
    #     ocr_text = ocr_gen_ai(Image.open(path))
    print("going to extract pdf")
    with ProcessPoolExecutor() as executor:
        doc_texts = list(executor.map(ocr_gen_ai, doc_pics))
    
    print("extracted_pdf")

    ocr_content = "\n".join(doc_texts)

    with open(f'{BASE_DIR}/skapp/ocr_text/{pdf_path.split("/")[-1].replace(".pdf", "")}.txt', "w", encoding="UTF-8") as oc:
        oc.write(ocr_content)

    return ocr_content

def ocr_gen_ai(path):
    # Choose a Gemini model.
    img = Image.open(path)
    model = genai.GenerativeModel(model_name="gemini-1.5-flash")

    prompt = """Extract text from the following image. Retain things in the image as their equivalent markdown in your response. Do not add any additional content to the content in the image (especially chatty texts)"""

    response = model.generate_content([prompt, "\n\n", img])

    return response.text
